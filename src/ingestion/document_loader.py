"""Document loading module for various file formats."""

import os
from typing import List, Dict
from pathlib import Path
import PyPDF2
import docx


class Document:
    """Represents a loaded document."""
    
    def __init__(self, content: str, metadata: Dict = None):
        self.content = content
        self.metadata = metadata or {}
    
    def __repr__(self):
        return f"Document(content_length={len(self.content)}, metadata={self.metadata})"


class DocumentLoader:
    """Load documents from various file formats."""
    
    SUPPORTED_FORMATS = {'.txt', '.pdf', '.docx'}
    
    def __init__(self, encoding: str = 'utf-8'):
        self.encoding = encoding
    
    def load_file(self, file_path: str) -> Document:
        """Load a single file and return a Document object."""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = path.suffix.lower()
        
        if extension not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported file format: {extension}")
        
        metadata = {
            'source': str(path),
            'filename': path.name,
            'extension': extension
        }
        
        if extension == '.txt':
            content = self._load_txt(path)
        elif extension == '.pdf':
            content = self._load_pdf(path)
        elif extension == '.docx':
            content = self._load_docx(path)
        else:
            raise ValueError(f"Unsupported format: {extension}")
        
        return Document(content=content, metadata=metadata)
    
    def load_directory(self, directory_path: str) -> List[Document]:
        """Load all supported documents from a directory."""
        documents = []
        path = Path(directory_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        for file_path in path.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_FORMATS:
                try:
                    doc = self.load_file(str(file_path))
                    documents.append(doc)
                    print(f"Loaded: {file_path.name}")
                except Exception as e:
                    print(f"Error loading {file_path.name}: {e}")
        
        return documents
    
    def _load_txt(self, path: Path) -> str:
        """Load text file."""
        with open(path, 'r', encoding=self.encoding) as f:
            return f.read()
    
    def _load_pdf(self, path: Path) -> str:
        """Load PDF file."""
        text = []
        with open(path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page in pdf_reader.pages:
                text.append(page.extract_text())
        return '\n'.join(text)
    
    def _load_docx(self, path: Path) -> str:
        """Load DOCX file."""
        doc = docx.Document(path)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
